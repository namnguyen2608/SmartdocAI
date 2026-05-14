# -*- coding: utf-8 -*-
"""
insert_diagrams.py
Render 2 diagram Mermaid (Self-RAG + Pipeline) ra PNG qua mermaid.ink,
sau đó chèn vào Word document với kích thước chuẩn A4 (width=14cm, canh giữa).

Dùng: python insert_diagrams.py
Output: diagrams_output.docx (mở bằng Word để xem)
"""

import base64
import sys
from io import BytesIO
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# ─────────────────────────────────────────────
# Mermaid source — Self-RAG (TD, compact)
# ─────────────────────────────────────────────
SELF_RAG_MERMAID = """%%{init: {'theme': 'base', 'themeVariables': {
  'primaryColor': '#EFF6FF',
  'primaryBorderColor': '#3B82F6',
  'secondaryColor': '#FEF9C3',
  'tertiaryColor': '#F0FDF4',
  'lineColor': '#374151',
  'fontSize': '15px'
}}}%%
flowchart TD
    A([Câu hỏi gốc]) --> B

    B["Bước 1 — Query Rewriting
    LLM viết lại thành 3 query variants"]

    B --> C["Bước 2 — Retrieve All Variants
    FAISS similarity_search × 3 queries
    Dedup theo 100 ký tự đầu content"]

    C --> D{"Bước 3 — Relevance Filter
    LLM chấm CÓ / KHÔNG
    cho từng document"}

    D -- CÓ --> E1["Giữ lại doc"]
    D -- KHÔNG --> E2["Loại bỏ doc"]
    E1 --> F
    E2 -.->|Nếu lọc hết giữ top-2| F

    F{"Bước 4 — Multi-hop Check
    LLM: needs_more_info?"}

    F -- Không --> G
    F -- Có --> F2["Sinh sub_questions
    Retrieve bổ sung thêm docs"]
    F2 --> G

    G["Bước 5 — Generate Answer
    format_context → Qwen2.5:7b
    RAG_PROMPT_TEMPLATE"]

    G --> H["Bước 6 — Answer Grading
    LLM → score · is_grounded
    has_hallucination · feedback"]

    H --> Z([Câu trả lời + Confidence Score + Sources])

    style B fill:#EFF6FF,stroke:#3B82F6
    style C fill:#EFF6FF,stroke:#3B82F6
    style D fill:#FEF9C3,stroke:#EAB308
    style F fill:#FCE7F3,stroke:#EC4899
    style G fill:#F0FDF4,stroke:#22C55E
    style H fill:#EDE9FE,stroke:#8B5CF6
"""

# ─────────────────────────────────────────────
# Mermaid source — Pipeline tổng thể (TD, compact)
# ─────────────────────────────────────────────
PIPELINE_MERMAID = """%%{init: {'theme': 'base', 'themeVariables': {
  'primaryColor': '#EFF6FF',
  'primaryBorderColor': '#3B82F6',
  'secondaryColor': '#FEF9C3',
  'tertiaryColor': '#F0FDF4',
  'lineColor': '#374151',
  'fontSize': '15px'
}}}%%
flowchart TD
    USER([Người dùng nhập câu hỏi])

    USER --> SR1

    subgraph SR1["Self-RAG — Tầng 1: Query Rewriting"]
        QR["rewrite_query
        LLM → 3 query variants"]
    end

    SR1 --> CORAG

    subgraph CORAG["Co-RAG — 3 Agent song song"]
        direction LR
        AG1["Agent 1
        Semantic
        FAISS MMR"]
        AG2["Agent 2
        Keyword
        BM25"]
        AG3["Agent 3
        Conceptual
        LLM sub-questions"]
    end

    CORAG --> CM

    subgraph CM["Consensus Merger"]
        CMN["Dedup + Vote Boost
        score = avg × (1 + (v-1)×0.15)
        CO_RAG_MIN_VOTES = 2"]
    end

    CM --> HY

    subgraph HY["Hybrid Search — RRF"]
        HYN["FAISS×0.6 + BM25×0.4
        RRF(d) = Σ 1/(60 + rank_i(d))"]
    end

    HY --> RE

    subgraph RE["CrossEncoder Reranking"]
        REN["ms-marco-MiniLM-L-6-v2
        Top-k = 5 docs"]
    end

    RE --> SR2

    subgraph SR2["Self-RAG — Tầng 2 & 3"]
        RF["Relevance Filter
        LLM → CÓ / KHÔNG / doc"]
        RF --> MH["Multi-hop + Generate
        Qwen2.5:7b"]
        MH --> AG["Answer Grading
        score · grounded · hallucination"]
    end

    SR2 --> CT

    subgraph CT["Citation Tracking"]
        CTN["Tên file + số trang
        từ metadata Document"]
    end

    CT --> OUT([Câu trả lời + Nguồn + Confidence → Streamlit])

    style SR1 fill:#FEF9C3,stroke:#EAB308
    style CORAG fill:#EFF6FF,stroke:#3B82F6
    style CM fill:#DBEAFE,stroke:#3B82F6
    style HY fill:#F0FDF4,stroke:#22C55E
    style RE fill:#DCFCE7,stroke:#16A34A
    style SR2 fill:#FEF9C3,stroke:#EAB308
    style CT fill:#EDE9FE,stroke:#8B5CF6
"""


def render_mermaid_png(mermaid_code: str, render_width: int = 1400) -> bytes:
    """Render Mermaid code → PNG bytes qua mermaid.ink API."""
    encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode()
    url = f"https://mermaid.ink/img/{encoded}?bgColor=white&width={render_width}"
    print(f"  Đang render từ mermaid.ink (width={render_width}px)...")
    response = requests.get(url, timeout=60)
    response.raise_for_status()
    return response.content


def set_para_spacing(para, before_pt: int = 0, after_pt: int = 6):
    """Thiết lập khoảng cách trước/sau đoạn."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before_pt * 20))
    spacing.set(qn("w:after"), str(after_pt * 20))
    pPr.append(spacing)


def insert_diagram(
    doc: Document,
    png_bytes: bytes,
    caption: str,
    width_cm: float = 14.0,
):
    """Chèn ảnh diagram + caption vào doc, canh giữa."""
    # ── Ảnh ──────────────────────────────────────
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(img_para, before_pt=6, after_pt=4)
    run = img_para.add_run()
    run.add_picture(BytesIO(png_bytes), width=Cm(width_cm))

    # ── Caption ───────────────────────────────────
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(cap_para, before_pt=2, after_pt=12)
    cap_run = cap_para.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    cap_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def main():
    out_path = Path("diagrams_output.docx")

    # ── Tạo document ──────────────────────────────
    doc = Document()

    # Lề A4 chuẩn: 2.5cm mỗi bên → usable width = 16cm
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Tiêu đề ───────────────────────────────────
    title = doc.add_heading("Hình minh họa kiến trúc SmartDoc AI", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Diagram 1: Self-RAG ───────────────────────
    print("[1/2] Render Self-RAG pipeline...")
    doc.add_heading("Hình 1. Self-RAG Pipeline (6 bước)", level=2)
    try:
        png1 = render_mermaid_png(SELF_RAG_MERMAID, render_width=1000)
        insert_diagram(
            doc, png1,
            caption="Hình 1. Self-RAG Pipeline — Query Rewriting → Retrieve → Relevance Filter → Multi-hop → Generate → Answer Grading",
            width_cm=14.0,
        )
        print("  ✓ Đã chèn Self-RAG diagram")
    except Exception as e:
        doc.add_paragraph(f"[Lỗi render Self-RAG: {e}]")
        print(f"  ✗ Lỗi: {e}")

    doc.add_page_break()

    # ── Diagram 2: Pipeline tổng thể ─────────────
    print("[2/2] Render Pipeline tổng thể...")
    doc.add_heading("Hình 2. Pipeline tổng thể SmartDoc AI", level=2)
    try:
        png2 = render_mermaid_png(PIPELINE_MERMAID, render_width=1000)
        insert_diagram(
            doc, png2,
            caption="Hình 2. Pipeline tổng thể — Co-RAG (3 agents) → Hybrid RRF → CrossEncoder → Self-RAG → Citation Tracking",
            width_cm=14.0,
        )
        print("  ✓ Đã chèn Pipeline diagram")
    except Exception as e:
        doc.add_paragraph(f"[Lỗi render Pipeline: {e}]")
        print(f"  ✗ Lỗi: {e}")

    # ── Lưu file ──────────────────────────────────
    doc.save(out_path)
    print(f"\n✓ Đã lưu: {out_path.resolve()}")
    print("  Mở file bằng Microsoft Word để kiểm tra.")


if __name__ == "__main__":
    main()
