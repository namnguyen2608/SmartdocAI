"""
SmartDocAI - Document Processor
Xử lý tài liệu PDF và DOCX: đọc, trích xuất văn bản, chia nhỏ (chunking)
"""

import os
import logging
from datetime import date
from typing import List, Optional

# from PyPDF2 import PdfReader
import pdfplumber
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

import config

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

def extract_text_from_pdf(file_path: str, source_name: Optional[str] = None) -> List[Document]:
    """
    Đọc file PDF và trích xuất văn bản từ từng trang.

    Args:
        file_path: Đường dẫn tới file PDF
        source_name: Tên hiển thị của file (mặc định dùng basename)

    Returns:
        Danh sách Document objects, mỗi object chứa nội dung 1 trang
        kèm metadata (source, page, file_type)

    Raises:
        FileNotFoundError: Khi file không tồn tại
        Exception: Khi file PDF bị hỏng hoặc không thể đọc
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Không tìm thấy file: {file_path}")

    documents = []
    file_name = source_name or os.path.basename(file_path)

    try:
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"Đang đọc file PDF '{file_name}' - {total_pages} trang")

            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    doc = Document(
                        page_content=text.strip(),
                        metadata={
                            "source": file_name,
                            "page": page_num + 1,
                            "total_pages": total_pages,
                            "file_type": "pdf",
                            "upload_date": date.today().isoformat(),
                        },
                    )
                    documents.append(doc)

        if not documents:
            logger.warning(
                f"File PDF '{file_name}' không chứa văn bản có thể trích xuất."
            )

        logger.info(
            f"Đã trích xuất {len(documents)} trang văn bản từ '{file_name}'"
        )

    except Exception as e:
        logger.error(f"Lỗi khi đọc file PDF '{file_name}': {str(e)}")
        raise Exception(
            f"Không thể đọc file PDF '{file_name}'. "
            f"File có thể bị hỏng hoặc được bảo vệ bằng mật khẩu. "
            f"Chi tiết: {str(e)}"
        )

    return documents

def extract_text_from_docx(file_path: str, source_name: Optional[str] = None) -> List[Document]:
    """
    Đọc file DOCX và trích xuất văn bản theo từng đoạn (paragraph).
    Các đoạn liên tiếp được ghép lại thành các "trang ảo" để phù hợp với
    pipeline chunking hiện tại.

    Args:
        file_path: Đường dẫn tới file DOCX
        source_name: Tên hiển thị của file (mặc định dùng basename)

    Returns:
        Danh sách Document objects kèm metadata (source, page, file_type)

    Raises:
        FileNotFoundError: Khi file không tồn tại
        Exception: Khi file DOCX bị hỏng hoặc không thể đọc
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Không tìm thấy file: {file_path}")

    file_name = source_name or os.path.basename(file_path)
    documents = []

    try:
        docx_doc = DocxDocument(file_path)

        # Thu thập toàn bộ văn bản từ paragraphs (bao gồm bảng)
        all_text_blocks: List[str] = []

        for para in docx_doc.paragraphs:
            text = para.text.strip()
            if text:
                all_text_blocks.append(text)

        for table in docx_doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    all_text_blocks.append(" | ".join(row_texts))

        if not all_text_blocks:
            logger.warning(
                f"File DOCX '{file_name}' không chứa văn bản có thể trích xuất."
            )
            return []

        # Gộp các block thành "trang ảo" ~1500 ký tự để phù hợp chunking
        VIRTUAL_PAGE_SIZE = 1500
        current_page_lines: List[str] = []
        current_len = 0
        virtual_page = 1

        for block in all_text_blocks:
            block_len = len(block)
            if current_len + block_len > VIRTUAL_PAGE_SIZE and current_page_lines:
                # Lưu trang ảo hiện tại
                page_text = "\n".join(current_page_lines)
                documents.append(
                    Document(
                        page_content=page_text,
                        metadata={
                            "source": file_name,
                            "page": virtual_page,
                            "total_pages": None,   # sẽ cập nhật sau
                            "file_type": "docx",
                            "upload_date": date.today().isoformat(),
                        },
                    )
                )
                virtual_page += 1
                current_page_lines = [block]
                current_len = block_len
            else:
                current_page_lines.append(block)
                current_len += block_len + 1  # +1 for newline

        # Trang cuối
        if current_page_lines:
            page_text = "\n".join(current_page_lines)
            documents.append(
                Document(
                    page_content=page_text,
                    metadata={
                        "source": file_name,
                        "page": virtual_page,
                        "total_pages": None,
                        "file_type": "docx",
                        "upload_date": date.today().isoformat(),
                    },
                )
            )

        # Cập nhật total_pages
        total_virtual_pages = len(documents)
        for doc in documents:
            doc.metadata["total_pages"] = total_virtual_pages

        logger.info(
            f"Đã trích xuất {len(all_text_blocks)} đoạn văn bản "
            f"({total_virtual_pages} trang ảo) từ '{file_name}'"
        )

    except Exception as e:
        logger.error(f"Lỗi khi đọc file DOCX '{file_name}': {str(e)}")
        raise Exception(
            f"Không thể đọc file DOCX '{file_name}'. "
            f"File có thể bị hỏng. "
            f"Chi tiết: {str(e)}"
        )

    return documents

def split_documents(
    documents: List[Document],
    chunk_size: Optional[int] = None,
    chunk_overlap: Optional[int] = None,
) -> List[Document]:
    """
    Chia nhỏ danh sách Document thành các chunks.

    Args:
        documents: Danh sách Document cần chia nhỏ
        chunk_size: Kích thước mỗi chunk (mặc định lấy từ config)
        chunk_overlap: Số ký tự overlap giữa các chunk (mặc định lấy từ config)

    Returns:
        Danh sách Document đã được chia nhỏ
    """
    if not documents:
        return []

    _chunk_size = chunk_size or config.CHUNK_SIZE
    _chunk_overlap = chunk_overlap or config.CHUNK_OVERLAP

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=_chunk_size,
        chunk_overlap=_chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    chunks = text_splitter.split_documents(documents)
    logger.info(
        f"Đã chia {len(documents)} trang thành {len(chunks)} chunks "
        f"(chunk_size={_chunk_size}, overlap={_chunk_overlap})"
    )

    return chunks

def process_uploaded_file(file_path: str, source_name: Optional[str] = None) -> List[Document]:
    """
    Pipeline hoàn chỉnh: Đọc file (PDF hoặc DOCX) → Trích xuất → Chunking.
    Tự động phát hiện định dạng từ phần mở rộng của file.

    Args:
        file_path: Đường dẫn tới file (PDF hoặc DOCX)
        source_name: Tên hiển thị tùy chọn

    Returns:
        Danh sách Document chunks đã xử lý

    Raises:
        ValueError: Khi định dạng file không được hỗ trợ
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".pdf":
        raw_documents = extract_text_from_pdf(file_path, source_name=source_name)
    elif ext == ".docx":
        raw_documents = extract_text_from_docx(file_path, source_name=source_name)
    else:
        raise ValueError(
            f"Định dạng file '{ext}' không được hỗ trợ. "
            f"Vui lòng sử dụng: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    if not raw_documents:
        return []

    chunks = split_documents(raw_documents)
    return chunks
